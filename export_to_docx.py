"""
Export PROJECT_OVERVIEW.txt to Word document
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Read the text file
with open('PROJECT_OVERVIEW.txt', 'r', encoding='utf-8') as f:
    content = f.read()

# Split into sections
sections = content.split('\n\n')

for section in sections:
    if not section.strip():
        continue

    lines = section.split('\n')
    first_line = lines[0].strip()

    # Check if it's a title (all caps or has === under it)
    is_title = False
    if len(lines) > 1 and lines[1].startswith('==='):
        is_title = True

    if is_title:
        # Add title
        heading = doc.add_heading(first_line, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Skip the === line
        remaining_lines = lines[2:] if len(lines) > 2 else []
    elif first_line.isupper() and len(first_line) > 3:
        # Section heading
        doc.add_heading(first_line, level=2)
        remaining_lines = lines[1:]
    else:
        remaining_lines = lines

    # Add remaining content
    if remaining_lines:
        text = '\n'.join(remaining_lines)
        if text.strip():
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

# Save document
doc.save('PROJECT_OVERVIEW.docx')
print("✅ Document exported to PROJECT_OVERVIEW.docx")
